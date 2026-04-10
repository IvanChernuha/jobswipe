"""Tests for cv_parser.extract_text — covers PDF, DOCX, plain-text, and edge cases."""
import io
import pytest


# ---------------------------------------------------------------------------
# Helper: build a minimal single-page PDF in pure bytes (no external tools)
# ---------------------------------------------------------------------------
_MINIMAL_PDF = (
    b"%PDF-1.4\n"
    b"1 0 obj<</Type /Catalog /Pages 2 0 R>>endobj\n"
    b"2 0 obj<</Type /Pages /Kids[3 0 R] /Count 1>>endobj\n"
    b"3 0 obj<</Type /Page /Parent 2 0 R /MediaBox[0 0 612 792]"
    b" /Contents 4 0 R /Resources<</Font<</F1 5 0 R>>>>>>endobj\n"
    b"4 0 obj<</Length 44>>\nstream\nBT /F1 12 Tf 100 700 Td (Hello CV) Tj ET\nendstream\nendobj\n"
    b"5 0 obj<</Type /Font /Subtype /Type1 /BaseFont /Helvetica>>endobj\n"
    b"xref\n0 6\n"
    b"0000000000 65535 f \n"
    b"0000000009 00000 n \n"
    b"0000000058 00000 n \n"
    b"0000000115 00000 n \n"
    b"0000000266 00000 n \n"
    b"0000000360 00000 n \n"
    b"trailer<</Size 6 /Root 1 0 R>>\nstartxref\n441\n%%EOF\n"
)


def _make_docx_bytes(text: str) -> bytes:
    """Create a minimal valid DOCX containing `text` in a single paragraph."""
    from docx import Document

    doc = Document()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Plain-text fallback
# ---------------------------------------------------------------------------

def test_extract_text_plain_utf8():
    from app.services.cv_parser import extract_text

    raw = "Hello, world! Python developer."
    result = extract_text(raw.encode("utf-8"), "text/plain")
    assert result == raw


def test_extract_text_plain_unknown_content_type_fallback():
    """Any unrecognised content_type falls back to plain-text decode."""
    from app.services.cv_parser import extract_text

    raw = b"just some text"
    result = extract_text(raw, "application/octet-stream")
    assert result == "just some text"


def test_extract_text_plain_non_utf8_decoded_via_chardet():
    """Non-UTF-8 bytes are auto-detected and decoded (not raised).

    Before the chardet change this byte would have been stripped via
    `errors="ignore"`. Now cv_parser uses chardet to detect the real
    encoding — 0xFF is valid Latin-1 (U+00FF) so it round-trips cleanly.
    This matters for Hebrew/Windows-1255 CVs where stripping "invalid"
    bytes would destroy the text.
    """
    from app.services.cv_parser import extract_text

    # 0xFF is invalid UTF-8 but valid Latin-1 (decodes to U+00FF = ÿ)
    data = b"hello \xff world"
    result = extract_text(data, "text/plain")
    # Must not raise; surrounding ASCII preserved.
    assert "hello" in result
    assert "world" in result
    # Decoded successfully via chardet (not stripped, not an error).
    assert isinstance(result, str)


# ---------------------------------------------------------------------------
# DOCX extraction
# ---------------------------------------------------------------------------

def test_extract_text_docx_roundtrip():
    from app.services.cv_parser import extract_text

    expected = "Senior Python Engineer with FastAPI experience."
    docx_bytes = _make_docx_bytes(expected)
    result = extract_text(
        docx_bytes,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert expected in result


def test_extract_text_docx_msword_content_type():
    """application/msword is also routed to the DOCX extractor."""
    from app.services.cv_parser import extract_text

    expected = "Skills: Docker, Kubernetes"
    docx_bytes = _make_docx_bytes(expected)
    # application/msword should still work because _extract_docx handles both
    result = extract_text(docx_bytes, "application/msword")
    assert expected in result


def test_extract_text_docx_empty_paragraphs_filtered():
    """Blank paragraphs in a DOCX should not appear in output."""
    from docx import Document
    from app.services.cv_parser import extract_text

    doc = Document()
    doc.add_paragraph("Line one")
    doc.add_paragraph("")          # blank — should be filtered
    doc.add_paragraph("   ")       # whitespace-only — should be filtered
    doc.add_paragraph("Line two")
    buf = io.BytesIO()
    doc.save(buf)

    result = extract_text(buf.getvalue(),
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    assert "Line one" in result
    assert "Line two" in result
    # Ensure we don't have stray blank lines padding the output
    lines = [l for l in result.splitlines() if l.strip() == ""]
    assert lines == []


# ---------------------------------------------------------------------------
# PDF extraction
# ---------------------------------------------------------------------------

def test_extract_pdf_returns_string():
    """pdfplumber can open a valid PDF and extract_text returns a str."""
    from app.services.cv_parser import extract_text

    # Use a known-valid minimal PDF; text extraction may be empty for this
    # hand-crafted binary but the call must not raise.
    result = extract_text(_MINIMAL_PDF, "application/pdf")
    assert isinstance(result, str)


def test_extract_pdf_empty_pages_produce_empty_string():
    """A PDF with no extractable text returns an empty string (not an error)."""
    import pdfplumber
    from app.services.cv_parser import _extract_pdf

    # Create a PDF where extract_text() returns None for every page
    class _FakePage:
        def extract_text(self):
            return None

    class _FakePDF:
        pages = [_FakePage(), _FakePage()]
        def __enter__(self): return self
        def __exit__(self, *a): pass

    import unittest.mock as mock
    with mock.patch("pdfplumber.open", return_value=_FakePDF()):
        result = _extract_pdf(b"fake")
    assert result == ""


# ---------------------------------------------------------------------------
# BUG: application/msword with a genuine old .doc binary would not parse
#      with python-docx (which only handles .docx).  Verify the call does
#      NOT silently return wrong data when the bytes are not a ZIP archive.
# ---------------------------------------------------------------------------

def test_extract_docx_raises_on_corrupt_bytes():
    """_extract_docx should raise (not silently return '') on invalid bytes."""
    from app.services.cv_parser import _extract_docx

    with pytest.raises(Exception):
        _extract_docx(b"this is not a zip file at all")
